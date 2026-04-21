from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^(\s*)[-*+]\s+(.*)$")
NUMBER_RE = re.compile(r"^(\s*)\d+[.)]\s+(.*)$")


def _list_style(base: str, level: int) -> str:
    level = max(0, min(level, 2))
    if level == 0:
        return base
    return f"{base} {level + 1}"


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # 基础正文样式（公文风）
    normal = doc.styles["Normal"]
    normal.font.name = "仿宋"
    normal.font.size = Pt(12)  # 小四
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    def style_heading(p, level: int) -> None:
        # 一级黑体三号，二级黑体小三，三级及以下黑体四号
        if level == 1:
            size = Pt(16)
        elif level == 2:
            size = Pt(15)
        else:
            size = Pt(14)
        for run in p.runs:
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.size = size
            run.bold = True
        pf = p.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5

    def style_body_paragraph(p) -> None:
        pf = p.paragraph_format
        pf.first_line_indent = Pt(24)  # 首行缩进2字符（按小四近似）
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        for run in p.runs:
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            run.font.size = Pt(12)

    def style_list_paragraph(p) -> None:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        for run in p.runs:
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            run.font.size = Pt(12)

    i = 0
    in_code = False
    code_lines: list[str] = []
    para_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal para_lines
        if not para_lines:
            return
        text = " ".join(s.strip() for s in para_lines if s.strip())
        if text:
            p = doc.add_paragraph(text)
            style_body_paragraph(p)
        para_lines = []

    def flush_code_block() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph("\n".join(code_lines))
        p.style = "No Spacing"
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        code_lines = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip("\n")
        stripped = line.strip()

        # fenced code block
        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                flush_code_block()
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # blank line
        if not stripped:
            flush_paragraph()
            i += 1
            continue

        # heading
        m = HEADING_RE.match(line)
        if m:
            flush_paragraph()
            level = min(len(m.group(1)), 4)
            p = doc.add_heading(m.group(2).strip(), level=level)
            style_heading(p, level)
            i += 1
            continue

        # bullet list
        m = BULLET_RE.match(line)
        if m:
            flush_paragraph()
            indent = len(m.group(1).replace("\t", "    "))
            level = indent // 2
            p = doc.add_paragraph(m.group(2).strip(), style=_list_style("List Bullet", level))
            style_list_paragraph(p)
            i += 1
            continue

        # numbered list
        m = NUMBER_RE.match(line)
        if m:
            flush_paragraph()
            indent = len(m.group(1).replace("\t", "    "))
            level = indent // 2
            p = doc.add_paragraph(m.group(2).strip(), style=_list_style("List Number", level))
            style_list_paragraph(p)
            i += 1
            continue

        # normal text paragraph block
        para_lines.append(line)
        i += 1

    flush_paragraph()
    if in_code:
        flush_code_block()

    doc.save(docx_path)


def main() -> int:
    if len(sys.argv) < 3:
        print("Usage: python convert_md_to_docx.py <input.md> <output.docx>")
        return 1
    md_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])
    convert_markdown_to_docx(md_path, docx_path)
    print(str(docx_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
